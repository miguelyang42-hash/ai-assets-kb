# -*- coding: utf-8 -*-
"""
V5.4 Word Report Generator（小蜂学掌定制版）
- 跨平台 (Win/Mac/Linux)
- 封面：小蜂学掌品牌封面（深海蓝标题块）
- 页眉：小蜂学掌 LOGO 文字 | 报告标题（双栏，深蓝分隔线）
- 页脚：© 小蜂学掌 www.xfxz123.com 版权声明（居中）
- 文件名前缀：小蜂学掌_
"""
from __future__ import annotations
import sys, json, os, re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 品牌常量 ──
BRAND_NAME    = "小蜂学掌"
BRAND_URL     = "www.xfxz123.com"
BRAND_FULL    = f"{BRAND_NAME}  |  精细化开发系统 V5.4"
BRAND_COPY    = f"© {BRAND_NAME}  {BRAND_URL}  版权所有，未经授权禁止转载"
NAVY_RGB      = RGBColor(0x1F, 0x4E, 0x79)   # 深海蓝 #1F4E79
GOLD_RGB      = RGBColor(0xF2, 0xC9, 0x4C)   # 品牌金

CHANNEL_LABEL = {
    "google":    "Google 搜索开发",
    "instagram": "Instagram 开发",
    "facebook":  "Facebook 开发",
    "linkedin":  "LinkedIn 开发",
    "tiktok":    "TikTok 开发",
    "manual":    "手工录入",
}


def safe_filename(name: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\r\n\t]+', '_', name).strip()
    return cleaned[:190]


def build_report_title(data: dict) -> str:
    v1 = data.get('v1_assessment', {})
    v2 = data.get('v2_intelligence', {})
    name_simple = v1.get('name_simple') or data.get('name_simple', 'Unknown')
    profile = v1.get('type', '潜在客户')
    country = v2.get('country') or v1.get('country', 'Unknown')
    rating = (v1.get('rating', 'B').split(' ')[0]) or 'B'
    return f"{name_simple} | {profile} | {country} | {rating}"


def _set_cell_bg(cell, hex_color: str):
    """给表格单元格设置背景色（ARGB）。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_brand_cover(doc: Document, title_text: str, channel: str, date_str: str):
    """插入品牌封面页。"""
    # 封面：深海蓝标题块（1 行 1 列表格，无边框）
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, '1F4E79')
    cell.width = Cm(16)

    # 封面第一行：品牌名
    p_brand = cell.paragraphs[0]
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_b = p_brand.add_run(BRAND_NAME)
    run_b.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run_b.font.size = Pt(28)
    run_b.bold = True

    # 封面第二行：网址
    p_url = cell.add_paragraph()
    p_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_u = p_url.add_run(BRAND_URL)
    run_u.font.color.rgb = GOLD_RGB
    run_u.font.size = Pt(13)

    # 封面第三行：分隔
    p_sep = cell.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = p_sep.add_run('─' * 36)
    r_sep.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r_sep.font.size = Pt(9)

    # 封面第四行：报告标题
    p_title = cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run('精细化客户开发报告')
    run_t.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run_t.font.size = Pt(18)
    run_t.bold = True

    # 封面第五行：客户名
    p_cname = cell.add_paragraph()
    p_cname.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cn = p_cname.add_run(title_text)
    run_cn.font.color.rgb = GOLD_RGB
    run_cn.font.size = Pt(13)

    # 封面第六行：渠道 + 日期
    p_meta = cell.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_m = p_meta.add_run(f"渠道：{CHANNEL_LABEL.get(channel, channel)}  |  {date_str}")
    run_m.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)
    run_m.font.size = Pt(10)

    doc.add_page_break()


def _add_header_footer(doc: Document, report_title: str):
    """向所有节添加页眉页脚。"""
    for section in doc.sections:
        section.different_first_page_header_footer = False

        # ── 页眉 ──
        header = section.header
        header.is_linked_to_previous = False
        # 清除默认段落
        for p in header.paragraphs:
            p.clear()
        if not header.paragraphs:
            header.add_paragraph()
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 左侧：品牌名（深蓝加粗）
        run_hl = hp.add_run(f"🐝 {BRAND_NAME}  |  {BRAND_URL}")
        run_hl.bold = True
        run_hl.font.color.rgb = NAVY_RGB
        run_hl.font.size = Pt(9)

        # Tab 隔开后右侧：报告简名
        hp.add_run('\t')
        run_hr = hp.add_run(f"精细化开发报告  |  V5.4")
        run_hr.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        run_hr.font.size = Pt(9)

        # 页眉下横线（paragraph border bottom）
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F4E79')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # ── 页脚 ──
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 版权声明
        run_fc = fp.add_run(BRAND_COPY)
        run_fc.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        run_fc.font.size = Pt(8)

        # 页码（分隔后）
        fp.add_run('    ')
        run_pg_label = fp.add_run('第 ')
        run_pg_label.font.size = Pt(8)
        run_pg_label.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        # 插入页码域
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_num = fp.add_run()
        run_num._r.append(fldChar1)
        run_num._r.append(instrText)
        run_num._r.append(fldChar2)
        run_num.font.size = Pt(8)
        run_pg_end = fp.add_run(' 页')
        run_pg_end.font.size = Pt(8)
        run_pg_end.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def create_pro_report(data_file: str, out_dir: str = None) -> str:
    with open(data_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    channel = data.get('channel', 'manual')
    date_str = datetime.now().strftime('%Y-%m-%d')

    # ── 封面 ──
    title_text = build_report_title(data)
    add_brand_cover(doc, title_text, channel, date_str)

    # ── 页眉页脚（在封面之后设置，覆盖所有节）──
    _add_header_footer(doc, title_text)

    # ── 渠道标签 ──
    p_chan = doc.add_paragraph()
    p_chan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_chan = p_chan.add_run(f"【信息来源渠道：{CHANNEL_LABEL.get(channel, channel)}】")
    run_chan.font.color.rgb = RGBColor(0, 112, 192)
    run_chan.bold = True

    p_sep = doc.add_paragraph('═' * 60)
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 一、客户类型判断与实力评估 ──
    doc.add_heading('一、客户类型判断与实力评估', level=1)
    v1 = data['v1_assessment']
    p1 = doc.add_paragraph()
    p1.add_run('客户类型（画像）：').bold = True
    p1.add_run(v1.get('type', ''))
    p1.add_run('\n判断依据：').bold = True
    p1.add_run(v1.get('basis', ''))
    p1.add_run('\n实力评分：').bold = True
    run_r = p1.add_run(v1.get('rating', ''))
    run_r.font.color.rgb = RGBColor(255, 0, 0)
    run_r.bold = True

    # ── 二、公司基本情报与社交资产 ──
    doc.add_heading('二、公司基本情报与社交资产', level=1)
    v2 = data['v2_intelligence']
    p2 = doc.add_paragraph()
    p2.add_run('公司英文全称：').bold = True
    p2.add_run(v2.get('company_full_name', ''))
    p2.add_run('\n公司英文简称：').bold = True
    p2.add_run(v1.get('name_simple', ''))
    p2.add_run('\n所在国家：').bold = True
    p2.add_run(v2.get('country', ''))
    p2.add_run('\n官网：').bold = True
    p2.add_run(v2.get('website', ''))
    p2.add_run('\n成立时间：').bold = True
    p2.add_run(v2.get('founded', ''))
    p2.add_run('\n总部地址：').bold = True
    p2.add_run(v2.get('hq', ''))
    p2.add_run('\n社交影响力：').bold = True
    p2.add_run(v2.get('social', ''))

    # ── 三、5 维深度分析 ──
    doc.add_heading('三、5 维深度分析 (Global Market Insight)', level=1)
    v3 = data['v3_insight']
    dims = [
        ('1. 工厂规模：', v3.get('factory_scale', '')),
        ('2. 市场定位：', v3.get('market_position', '')),
        ('3. 社交活跃度：', v3.get('social_activity', '')),
        ('4. 采购逻辑：', v3.get('sourcing_logic', '')),
        ('5. 技术偏好：', v3.get('tech_preference', '')),
    ]
    for label, content in dims:
        p = doc.add_paragraph(style='List Number')
        p.add_run(label).bold = True
        p.add_run(content)

    # ── 四、关键联系人深度图谱 ──
    doc.add_heading('四、关键联系人深度图谱 (Key Contacts Map)', level=1)
    for i, m in enumerate(data['v4_contacts'], 1):
        p_m = doc.add_paragraph()
        run_h = p_m.add_run(f'★ 关键联系人 {i}')
        run_h.bold = True
        run_h.font.color.rgb = RGBColor(0, 70, 140)
        for label, key in [
            ('姓名', 'name'), ('职位', 'title'), ('邮箱地址', 'email'),
            ('LinkedIn地址', 'linkedin'), ('WhatsApp/Tel', 'whatsapp'),
            ('推荐第一触点', 'touchpoint'), ('互动策略', 'strategy'),
        ]:
            p_info = doc.add_paragraph(style='List Bullet')
            p_info.add_run(f'{label}：').bold = True
            p_info.add_run(str(m.get(key, '')))

    # ── 五、谈判策略 ──
    doc.add_heading('五、商务谈判专家策略分析', level=1)
    v5 = data['v5_negotiation']
    p5 = doc.add_paragraph()
    p5.add_run('局面判断：').bold = True
    p5.add_run(v5.get('situation', ''))
    p5.add_run('\n核心卡点：').bold = True
    p5.add_run(v5.get('pain_point', ''))
    p5.add_run('\n应对方案：').bold = True
    p5.add_run(v5.get('solution', ''))

    # ── 六、双语开发信 ──
    doc.add_heading('六、定制化开发信模板 (Bilingual)', level=1)
    p_zh = doc.add_paragraph()
    p_zh.add_run('【中文邮件模板】').bold = True
    doc.add_paragraph(data['v6_templates'].get('chinese', ''))
    p_en = doc.add_paragraph()
    p_en.add_run('【English Official Outreach】').bold = True
    doc.add_paragraph(data['v6_templates'].get('english', ''))

    # ── 七、产品业务模式 ──
    doc.add_heading('七、产品与业务模式分析', level=1)
    v7 = data['v7_product_analysis']
    p7 = doc.add_paragraph()
    p7.add_run('价格定位：').bold = True
    p7.add_run(v7.get('pricing', ''))
    p7.add_run('\n供应链偏好：').bold = True
    p7.add_run(v7.get('preference', ''))

    # ── 八、落地建议 ──
    doc.add_heading('八、业务落地建议', level=1)
    for sug in data.get('v8_suggestions', []):
        doc.add_paragraph(sug, style='List Bullet')

    # ── 保存（文件名前缀：小蜂学掌_）──
    safe_title = safe_filename(title_text)
    save_dir = Path(out_dir) if out_dir else Path(os.environ.get('V5_OUT_DIR', '.'))
    save_dir.mkdir(parents=True, exist_ok=True)
    save_path = save_dir / f"小蜂学掌_{safe_title}.docx"
    doc.save(str(save_path))
    abs_path = str(save_path.resolve())
    print(f"WORD_READY: {abs_path}")
    return abs_path


if __name__ == "__main__":
    if len(sys.argv) > 1:
        out_dir = sys.argv[2] if len(sys.argv) > 2 else None
        create_pro_report(sys.argv[1], out_dir)
